from fastapi import APIRouter, Depends, HTTPException, Query
from sqlalchemy.orm import Session
import math
import re
from app.db.session import get_db
from app.core.security import get_current_user, require_roles
from app.models.all_models import (
    Project, ProjectMember, UserRole, User, AuditLog, AuditAction, ProjectStatus
)
from app.schemas.api_schemas import ProjectCreate, ProjectOut, ProjectUpdate, ProjectMemberAdd, ProjectMemberOut, PaginatedResponse

router = APIRouter(prefix="/projects", tags=["projects"])


def _serialize(p: Project, db: Session) -> ProjectOut:
    return ProjectOut(
        id=p.id,
        name=p.name,
        description=p.description,
        status=p.status.value,
        owner_id=p.owner_id,
        submission_destinations=p.submission_destinations or [],
        created_at=p.created_at,
        member_count=len(p.members),
        job_count=len(p.jobs),
    )


def _can_access(user: User, project: Project) -> bool:
    user_roles = {r.role.value for r in user.roles}
    if "org_admin" in user_roles:
        return True
    return any(m.user_id == user.id for m in project.members)


@router.get("", response_model=PaginatedResponse)
def list_projects(
    page: int = Query(1, ge=1),
    page_size: int = Query(20, ge=1, le=100),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    user_roles = {r.role.value for r in current_user.roles}
    q = db.query(Project).filter(Project.deleted_at == None)
    if "org_admin" not in user_roles:
        q = q.join(ProjectMember).filter(ProjectMember.user_id == current_user.id)
    total = q.count()
    projects = q.offset((page - 1) * page_size).limit(page_size).all()
    return PaginatedResponse(
        items=[_serialize(p, db) for p in projects],
        total=total, page=page, page_size=page_size,
        pages=math.ceil(total / page_size),
    )


@router.post("", response_model=ProjectOut, status_code=201)
def create_project(
    payload: ProjectCreate,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_roles("org_admin", "project_admin")),
):
    project = Project(
        name=payload.name,
        description=payload.description,
        submission_destinations=payload.submission_destinations,
        owner_id=current_user.id,
    )
    db.add(project)
    db.flush()
    db.add(ProjectMember(project_id=project.id, user_id=current_user.id, role=UserRole.PROJECT_ADMIN))
    db.add(AuditLog(
        user_id=current_user.id, project_id=project.id,
        action=AuditAction.PROJECT_CREATED,
        after_value={"name": payload.name},
    ))
    db.commit()
    db.refresh(project)
    return _serialize(project, db)


@router.get("/{project_id}", response_model=ProjectOut)
def get_project(
    project_id: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    project = db.query(Project).filter(Project.id == project_id, Project.deleted_at == None).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    if not _can_access(current_user, project):
        raise HTTPException(status_code=403, detail="Access denied")
    return _serialize(project, db)


@router.patch("/{project_id}", response_model=ProjectOut)
def update_project(
    project_id: str,
    payload: ProjectUpdate,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    project = db.query(Project).filter(Project.id == project_id, Project.deleted_at == None).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    user_roles = {r.role.value for r in current_user.roles}
    is_project_admin = any(
        m.user_id == current_user.id and m.role in (UserRole.PROJECT_ADMIN,)
        for m in project.members
    )
    if "org_admin" not in user_roles and not is_project_admin:
        raise HTTPException(status_code=403, detail="Access denied")

    before_status = project.status.value
    if payload.name: project.name = payload.name
    if payload.description is not None: project.description = payload.description
    if payload.submission_destinations is not None:
        project.submission_destinations = payload.submission_destinations
    if payload.status:
        project.status = ProjectStatus(payload.status)
        db.add(AuditLog(
            user_id=current_user.id, project_id=project.id,
            action=AuditAction.PROJECT_STATUS_CHANGED,
            before_value={"status": before_status},
            after_value={"status": payload.status},
        ))
    db.commit()
    db.refresh(project)
    return _serialize(project, db)


@router.get("/{project_id}/members", response_model=list[ProjectMemberOut])
def list_members(
    project_id: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    project = db.query(Project).filter(Project.id == project_id, Project.deleted_at == None).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    if not _can_access(current_user, project):
        raise HTTPException(status_code=403, detail="Access denied")
    return [
        ProjectMemberOut(
            user_id=m.user_id, role=m.role.value,
            full_name=m.user.full_name, email=m.user.email,
            created_at=m.created_at,
        )
        for m in project.members
    ]


@router.post("/{project_id}/members", status_code=201)
def add_member(
    project_id: str,
    payload: ProjectMemberAdd,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_roles("org_admin", "project_admin")),
):
    project = db.query(Project).filter(Project.id == project_id).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    user = db.query(User).filter(User.id == payload.user_id).first()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    existing = db.query(ProjectMember).filter(
        ProjectMember.project_id == project_id,
        ProjectMember.user_id == payload.user_id,
    ).first()
    if existing:
        existing.role = UserRole(payload.role)
    else:
        db.add(ProjectMember(project_id=project_id, user_id=payload.user_id, role=UserRole(payload.role)))
    db.commit()
    return {"message": "Member added"}


@router.delete("/{project_id}/members/{user_id}", status_code=204)
def remove_member(
    project_id: str,
    user_id: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_roles("org_admin", "project_admin")),
):
    member = db.query(ProjectMember).filter(
        ProjectMember.project_id == project_id, ProjectMember.user_id == user_id
    ).first()
    if not member:
        raise HTTPException(status_code=404, detail="Member not found")
    db.delete(member)
    db.commit()


@router.delete("/{project_id}", status_code=204)
def delete_project(
    project_id: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_roles("org_admin")),
):
    """Soft-delete a project. Only org admins can delete projects."""
    from datetime import datetime, timezone
    project = db.query(Project).filter(Project.id == project_id, Project.deleted_at == None).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    project.deleted_at = datetime.now(timezone.utc)
    db.add(AuditLog(
        user_id=current_user.id, project_id=project_id,
        action=AuditAction.PROJECT_DELETED,
        before_value={"name": project.name},
    ))
    db.commit()



@router.get("/{project_id}/export-preview")
def export_preview(
    project_id: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    Returns approved records for preview in the UI before downloading.
    Groups records by source so the user can browse them.
    """
    from app.models.all_models import Source, ExtractionJob, ExtractedRecord, ReviewStatus, SourceStatus

    project = db.query(Project).filter(Project.id == project_id).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    approved_sources = db.query(Source).filter(
        Source.project_id == project_id,
        Source.status == SourceStatus.APPROVED,
    ).order_by(Source.name).all()

    result = []
    for source in approved_sources:
        jobs = db.query(ExtractionJob).filter(
            ExtractionJob.source_id == source.id
        ).all()
        source_records = []
        for job in jobs:
            recs = db.query(ExtractedRecord).filter(
                ExtractedRecord.job_id == job.id,
                ExtractedRecord.review_status == ReviewStatus.APPROVED,
            ).all()
            for r in recs:
                ef = r.extracted_fields or {}
                source_records.append({
                    "record_id":      str(r.id),
                    "canonical_name": ef.get("canonical_name") or ef.get("company_name") or source.name,
                    "company_name":   ef.get("company_name") or source.name,
                    "is_submitted":   bool(getattr(r, "is_submitted", False)),
                    "fields":         ef,
                })
        if source_records:
            result.append({
                "source_id":    str(source.id),
                "source_name":  source.name,
                "approved_at":  source.approved_at.isoformat() if source.approved_at else None,
                "record_count": len(source_records),
                "records":      source_records,
            })

    return {
        "project":        project.name,
        "project_folder": project.name,
        "total_sources":  len(result),
        "total_records":  sum(s["record_count"] for s in result),
        "sources":        result,
    }

@router.get("/{project_id}/export-debug")
def export_debug(
    project_id: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    Debug endpoint — returns JSON showing exactly what the export query finds.
    Use this to verify data exists before trying to download the ZIP.
    """
    from app.models.all_models import Source, ExtractionJob, ExtractedRecord, ReviewStatus, SourceStatus
    project = db.query(Project).filter(Project.id == project_id).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    sources = db.query(Source).filter(Source.project_id == project_id).all()
    result = {
        "project": project.name,
        "total_sources": len(sources),
        "sources": []
    }
    for source in sources:
        jobs = db.query(ExtractionJob).filter(ExtractionJob.source_id == source.id).all()
        source_info = {
            "name": source.name,
            "status": source.status.value,
            "jobs": len(jobs),
            "records_per_job": []
        }
        for job in jobs:
            recs = db.query(ExtractedRecord).filter(ExtractedRecord.job_id == job.id).all()
            approved = [r for r in recs if str(getattr(r.review_status, 'value', r.review_status)) == 'approved']
            source_info["records_per_job"].append({
                "job_id": str(job.id),
                "total_records": len(recs),
                "approved_records": len(approved),
                "has_extracted_fields": sum(1 for r in recs if r.extracted_fields)
            })
        result["sources"].append(source_info)
    return result


@router.get("/{project_id}/export")
def export_project(
    project_id: str,
    status: str = "all",
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    Download all records from a project as a ZIP.
    
    ?status=all       → every record from every source
    ?status=approved  → only sources with status=approved, approved records only
    
    Logic:
      Source has no deleted_at.
      Records link to sources through ExtractionJob (not directly).
      Path: Source → ExtractionJob.source_id → ExtractedRecord.job_id
    """
    import io, zipfile, json as _json
    from datetime import datetime, timezone
    from fastapi.responses import StreamingResponse
    from app.models.all_models import Source, ExtractionJob, ExtractedRecord, ReviewStatus

    # 1. Verify project exists
    project = db.query(Project).filter(Project.id == project_id).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    # 2. Get sources — filter to approved only if requested
    source_q = db.query(Source).filter(Source.project_id == project_id)
    if status == "approved":
        from app.models.all_models import SourceStatus as SS
        source_q = source_q.filter(Source.status == SS.APPROVED)
    sources = source_q.all()

    if not sources:
        raise HTTPException(
            status_code=404,
            detail="No approved sources found in this project. Approve sources first." if status == "approved"
                   else "No sources found in this project."
        )

    # 3. Collect records through job_id path (Source has no direct records)
    all_records = []
    for source in sources:
        # Find all extraction jobs for this source
        jobs = db.query(ExtractionJob).filter(
            ExtractionJob.source_id == source.id
        ).all()

        for job in jobs:
            # Back to ORM — safe now that LLMVerdict enum includes SKIPPED + FLAGGED
            rec_q = db.query(ExtractedRecord).filter(
                ExtractedRecord.job_id == job.id
            )
            if status == "approved":
                rec_q = rec_q.filter(
                    ExtractedRecord.review_status == ReviewStatus.APPROVED
                )

            for r in rec_q.all():
                try:
                    ef = dict(r.extracted_fields) if isinstance(r.extracted_fields, dict) else (r.extracted_fields or {})
                    review_str = r.review_status.value if hasattr(r.review_status, "value") else str(r.review_status)
                    ef["_"] = {
                        "source_name":    source.name,
                        "source_status":  source.status.value,
                        "review_status":  review_str,
                        "is_submitted":   bool(getattr(r, "is_submitted", False)),
                        "is_schema_valid": bool(getattr(r, "is_schema_valid", False)),
                        "exported_at":    datetime.now(timezone.utc).isoformat(),
                    }
                    cn = ef.get("canonical_name") or ef.get("company_name") or source.name
                    all_records.append({
                        "source_name":    source.name,
                        "canonical_name": cn,
                        "data":           ef,
                    })
                except Exception:
                    continue

    if not all_records:
        raise HTTPException(
            status_code=404,
            detail="No records found. Upload data to sources and approve records first."
        )

    # 4. Build ZIP
    ts = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    readme_lines = [
        f"# {project.name} —  Export",
        f"",
        f"Exported: {ts}",
        f"Filter:   {status}",
        f"Records:  {len(all_records)}",
        f"Sources:  {len(sources)}",
        f"",
        f"## Files",
        f"- records/   one JSON per company (named by canonical_name)",
        f"- combined.json   all records in one array",
        f"- README.md   this file",
        f"",
        f"## Sources included",
    ]
    for s in sources:
        count = sum(1 for r in all_records if r["source_name"] == s.name)
        if count:
            readme_lines.append(f"- {s.name}: {count} record(s) [{s.status.value}]")

    # ZIP folder = project name (e.g. "Critical Materials Intelligence/")
    # Files sit directly inside: albemarle-corporation.json, combined.json, README.md
    project_folder = project.name  # exact project name — no transformation
    project_slug   = re.sub(r"[^a-z0-9_]", "_", project.name.lower())[:40]

    buf = io.BytesIO()
    seen: dict = {}
    combined = []

    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr(f"{project_folder}/README.md", "\n".join(readme_lines))

        for item in all_records:
            # canonical_name is the SOP tracking key — never transform it
            cn = str(item["canonical_name"]).strip()
            cn = cn.replace("/", "-").replace("\\", "-").replace(":", "-").replace("\n", "") or "record"
            if cn in seen:
                seen[cn] += 1
                cn = f"{cn}_{seen[cn]}"
            else:
                seen[cn] = 0

            record_json = item["data"]
            combined.append(record_json)
            # File lives directly inside the project folder
            zf.writestr(
                f"{project_folder}/{cn}.json",
                _json.dumps(record_json, indent=2, ensure_ascii=False, default=str)
            )

        zf.writestr(
            f"{project_folder}/combined.json",
            _json.dumps(combined, indent=2, ensure_ascii=False, default=str)
        )

    buf.seek(0)
    zip_filename = f"{project_slug}_{status}_{datetime.now().strftime('%Y%m%d')}.zip"

    return StreamingResponse(
        buf,
        media_type="application/zip",
        headers={"Content-Disposition": f'attachment; filename="{zip_filename}"'},
    )


@router.get("/{project_id}/export-package")
def export_project_package(
    project_id: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    Full project download package:
    - A Word document with one cover page per source (extractor, reviewer, stats, dates)
    - All JSON records in records/ folder
    - combined.json with every record
    - README.md summary

    This is the single ZIP the client downloads to get everything at once.
    """
    import io, zipfile, json as _json
    from datetime import datetime, timezone
    from fastapi.responses import StreamingResponse
    from app.models.all_models import Source, ExtractedRecord, ReviewStatus, User as UserModel
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    project = db.query(Project).filter(
        Project.id == project_id, Project.deleted_at == None
    ).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    sources = db.query(Source).filter(
        Source.project_id == project_id,
    ).order_by(Source.name).all()

    ts_label = datetime.now(timezone.utc).strftime("%d %B %Y, %H:%M UTC")
    ts_file  = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M")

    # ── Resolve user names ────────────────────────────────────────────────────
    user_cache: dict[str, str] = {}
    def get_name(uid: str | None) -> str:
        if not uid: return "Unassigned"
        if uid not in user_cache:
            u = db.query(UserModel).filter(UserModel.id == uid).first()
            user_cache[uid] = u.full_name or u.email if u else uid[:8]
        return user_cache[uid]

    # ── Collect all records (via ExtractionJob — no direct source_id on records) ──
    from app.models.all_models import ExtractionJob as EJ
    all_records = []
    for source in sources:
        job_ids = [j.id for j in db.query(EJ).filter(EJ.source_id == str(source.id)).all()]
        if not job_ids:
            continue
        recs = db.query(ExtractedRecord).filter(
            ExtractedRecord.job_id.in_(job_ids)
        ).all()
        for r in recs:
            d = dict(r.extracted_fields or {})
            d["_"] = {
                "source": source.name,
                "review_status": r.review_status.value if hasattr(r.review_status, "value") else r.review_status,
                "is_schema_valid": r.is_schema_valid,
                "exported_at": datetime.now(timezone.utc).isoformat(),
            }
            all_records.append(d)

    # ── Build Word cover-page document ────────────────────────────────────────
    doc = DocxDocument()

    # Page margins
    section = doc.sections[0]
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.0)

    def set_cell_bg(cell, hex_color: str):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  hex_color)
        tcPr.append(shd)

    def add_run(para, text, bold=False, italic=False, size=11, color=None, font="Calibri"):
        run = para.add_run(text)
        run.bold   = bold
        run.italic = italic
        run.font.name = font
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor.from_string(color)
        return run

    def cell_para(cell, text, bold=False, size=10, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
        cell.paragraphs[0].clear()
        p = cell.paragraphs[0]
        p.alignment = align
        add_run(p, text, bold=bold, size=size, color=color)

    # ── PROJECT TITLE PAGE ────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    add_run(p, project.name.upper(), bold=True, size=24, color="1B3A6B")

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p2, "Data Extraction & Review Package", size=14, color="2563EB")

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(12)
    add_run(p3, f"Generated: {ts_label}", size=10, color="64748B", italic=True)

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p4, f"Total Sources: {len(sources)}  ·  Total Records: {len(all_records)}", size=11, color="475569")

    # Summary table
    doc.add_paragraph().paragraph_format.space_before = Pt(24)
    stats_by_status: dict[str, int] = {}
    for s in sources:
        v = s.status.value if hasattr(s.status, "value") else s.status
        stats_by_status[v] = stats_by_status.get(v, 0) + 1

    tbl = doc.add_table(rows=1 + len(stats_by_status), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    hdr = tbl.rows[0]
    cell_para(hdr.cells[0], "Source Status", bold=True, size=10, color="FFFFFF")
    cell_para(hdr.cells[1], "Count",         bold=True, size=10, color="FFFFFF", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_bg(hdr.cells[0], "1B3A6B")
    set_cell_bg(hdr.cells[1], "1B3A6B")
    for i, (status, count) in enumerate(sorted(stats_by_status.items())):
        row = tbl.rows[i + 1]
        cell_para(row.cells[0], status.replace("_", " ").title(), size=10)
        cell_para(row.cells[1], str(count), size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        bg = "ECFDF5" if status == "approved" else "FFFBEB" if "review" in status else "F8FAFC"
        set_cell_bg(row.cells[0], bg)
        set_cell_bg(row.cells[1], bg)

    doc.add_page_break()

    # ── ONE COVER PAGE PER SOURCE ─────────────────────────────────────────────
    for s_idx, source in enumerate(sources):
        records_for_source = [r for r in all_records if r.get("_", {}).get("source") == source.name]
        approved_count = sum(1 for r in records_for_source if r.get("_", {}).get("review_status") == "approved")
        total_count    = len(records_for_source)

        extractor_name = get_name(source.assigned_extractor_id)
        reviewer_name  = get_name(source.assigned_reviewer_id)
        status_val     = source.status.value if hasattr(source.status, "value") else source.status
        status_label   = status_val.replace("_", " ").title()
        status_color   = "059669" if status_val == "approved" else "D97706" if "review" in status_val else "64748B"

        # Source heading
        ph = doc.add_paragraph()
        ph.paragraph_format.space_before = Pt(6)
        add_run(ph, source.name, bold=True, size=16, color="1B3A6B")

        # Status line
        ps = doc.add_paragraph()
        add_run(ps, f"Status: ", bold=True, size=10, color="374151")
        add_run(ps, status_label, bold=True, size=10, color=status_color)
        if source.website_url:
            add_run(ps, f"  ·  {source.website_url}", size=9, color="2563EB")

        # Details table
        details = [
            ("Extractor",          extractor_name),
            ("Reviewer",           reviewer_name),
            ("Schema",             source.schema_name or "—"),
            ("Total Records",      str(total_count)),
            ("Approved Records",   str(approved_count)),
            ("Pending / Rejected", str(total_count - approved_count)),
            ("Extraction Started", source.extraction_started_at.strftime("%d %b %Y %H:%M") if source.extraction_started_at else "—"),
            ("Approved At",        source.approved_at.strftime("%d %b %Y %H:%M") if source.approved_at else "—"),
            ("Description",        source.description or "—"),
        ]

        t = doc.add_table(rows=len(details), cols=2)
        t.style = "Table Grid"
        for i, (lbl, val) in enumerate(details):
            row = t.rows[i]
            cell_para(row.cells[0], lbl, bold=True, size=9, color="374151")
            cell_para(row.cells[1], val, size=9)
            set_cell_bg(row.cells[0], "EFF6FF")
            set_cell_bg(row.cells[1], "FFFFFF" if i % 2 == 0 else "F8FAFC")

        # Records summary (first 5 company names)
        if records_for_source:
            doc.add_paragraph().paragraph_format.space_before = Pt(8)
            pr = doc.add_paragraph()
            add_run(pr, "Sample Records: ", bold=True, size=9, color="374151")
            names = [r.get("company_name") or r.get("material_name") or "—" for r in records_for_source[:5]]
            add_run(pr, ", ".join(str(n) for n in names) + (f" … and {total_count-5} more" if total_count > 5 else ""), size=9, color="64748B")

        # Horizontal rule spacer before next source
        if s_idx < len(sources) - 1:
            doc.add_paragraph().paragraph_format.space_before = Pt(4)
            sep = doc.add_paragraph()
            sep.paragraph_format.space_after = Pt(4)
            add_run(sep, "─" * 90, size=7, color="E2E8F0")

    # ── Pack everything into ZIP ──────────────────────────────────────────────
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:

        # Word cover document
        word_buf = io.BytesIO()
        doc.save(word_buf)
        word_buf.seek(0)
        project_slug = re.sub(r"[^a-z0-9_]", "_", project.name.lower())[:40]
        zf.writestr(f"{project_slug}_review_package.docx", word_buf.read())

        # Combined JSON
        zf.writestr("combined_records.json",
                    _json.dumps(all_records, indent=2, ensure_ascii=False, default=str))

        # Individual JSON records — use canonical_name exactly as stored per SOP
        seen: dict[str, int] = {}
        for record in all_records:
            cn = str(record.get("canonical_name") or record.get("company_name") or "record").strip()
            cn = cn.replace("/", "-").replace("\\", "-").replace(":", "-").replace("\n", "") or "record"
            if cn in seen:
                seen[cn] += 1; cn = f"{cn}_{seen[cn]}"
            else:
                seen[cn] = 0
            zf.writestr(f"records/{cn}.json",
                        _json.dumps(record, indent=2, ensure_ascii=False, default=str))

        # README
        readme = (
            f"# {project.name} — Review Package\n\n"
            f"Generated: {ts_label}\n"
            f"Sources: {len(sources)} | Records: {len(all_records)}\n\n"
            f"## Contents\n"
            f"- `{project_slug}_review_package.docx` — Cover pages for every source (extractor, reviewer, dates, stats)\n"
            f"- `records/` — One JSON file per company record\n"
            f"- `combined_records.json` — All records in one array\n"
            f"- `README.md` — This file\n\n"
            f"## Sources ({len(sources)})\n"
        )
        for source in sources:
            sv = source.status.value if hasattr(source.status, "value") else source.status
            readme += f"- {source.name} [{sv}] — Extractor: {get_name(source.assigned_extractor_id)}, Reviewer: {get_name(source.assigned_reviewer_id)}\n"
        zf.writestr("README.md", readme)

    buf.seek(0)
    fname = f"{project_slug}_package_{ts_file}.zip"
    return StreamingResponse(
        buf,
        media_type="application/zip",
        headers={"Content-Disposition": f"attachment; filename={fname}"},
    )
